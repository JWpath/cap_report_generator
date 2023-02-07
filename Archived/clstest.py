import docx
import docx.shared
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re

class Section:
    """
    a Section is a text structured as a Header followed by choices
    
    Consider a list of strings (docx.paragraphs) structured as follows : 
    
    head
    Dict of choices : [choicea0, choice1, choice2, ..., choiceN] 
        ___ choice0 : 1
        ___ choice1 : 2
        ___ choice2 : 3 
        ___ choice3 : 4
        ___ choice4 : 5
        
    end of defined by either empty string, or another head or a optionhead
    
    Example : Margins = Section(headtext, choice list)
     
    """

    def __init__(self, head, choice, optionhead = None, optionchoice = None):
        self.head = head
        self.choice = choice #list
        
        ### Trying to implement optiobnheads
        self.optionhead = optionhead 
        self.optionchoice = optionchoice
        
        
    @classmethod
    #Create section without instantiating a section object first
    def get_section(cls, paragraphs_list) :
        
        choice = {}
        for i, paragraph in enumerate(paragraphs_list) :
            
            #if matches header regex, set as header [SOLVED]
            if headmatch := re.search(r"^((?:[A-Z]?[a-z0-9 ]+)+).*$", paragraphs_list[i].text) :
                head = headmatch.group(1)
            #if matches choice regex, append to choice list [SOLVED]
            elif choicematch := re.search(r"^_+([a-zA-Z0-9 ,.'()/\\?!]+):?.*$", paragraphs_list[i].text) : 
                print(i, choicematch.group(1), type(choicematch.group(1)))
                choice.update({choicematch.group(1):i})
                #choice.append(f"{str(i)} : {choicematch.group(1)}")
            # break if next line is not compatible with choice format
            
            if not re.search(r"^_+([a-zA-Z0-9 ]+):?.*$", paragraphs_list[i+1].text) :
                break
        return cls(head, choice)
            

            # break if next line is not compatible with choice format
            #if not re.search(r"^_+([a-zA-Z0-9 ]+):?.*$", paragraphs_list[i+1].text) :
                #break
    
def main():
    cap = docx.Document("tumorextent_section_test.docx")
    caplist = cap.paragraphs
    
    capsection = Section.get_section(caplist)
    
    print(capsection.head)
    print(capsection.choice)
    
    
    for values in capsection.choice.values():
         print(values)

#Extract header
#Extract choice list


def extract_template (d, start = None, end = None):
    """ 
        Args : 
            d : a `docx.Document()` object
            start : string at which extraction begins
            end : string at which extraction ends
        Return : 
            `d.paragraphs` list from specified from `start` to `end`.
    """
    if start == None : 
        start = input("Initiate text extraction at : ")
    if end == None : 
        end = input("Terminate text extraction at : ")
    
    for i, paragraphs in enumerate(d.paragraphs) :     
        if d.paragraphs[i].text.startswith(str(start)) :  #and is char style.bold == TRUE and is capitalized
            start_index = i
            #print(f"start_index is : {start_index}")
        elif d.paragraphs[i].text.startswith(str(end)) : #and is char style.bold == TRUE and is capitalized
            end_index = i
            #print(f"end_index is : {end_index}")
            break   
    return d.paragraphs[start_index:end_index]

def new_paragraph_style (document, stylename = None, font_name = "Calibri", 
              font_size = 11, 
              font_bold = False, font_italic = False, font_underline = False, 
              color = docx.shared.RGBColor(0, 0, 0),
              before_spacing = 0,
              after_spacing = 0,
              keep_together = True,
              line_spacing = 1.0
              ):
    """
    Args : 
        document : a `docx.Document()` object
        stylename : name of newstyle to be added
    
    Return : 
        a new `style` object with specified font and paragraph_format parameters
    """
    #If stylename argument not defined, prompt user for new stylename
    if stylename == None : 
        stylename = input("New style name : ")
    
    #function will set new style if style does not already exist    
    if stylename in (document.styles) : 
        raise ValueError("Style already exists")
    
    else : 
        #adjust newstyle.font parameters
        newstyle = (document.styles).add_style(stylename, WD_STYLE_TYPE.PARAGRAPH)
        newstyle.font.name = font_name
        newstyle.font.size = Pt(font_size)
        newstyle.font.bold = font_bold
        newstyle.font.italic = font_italic
        newstyle.font.underline = font_underline
        newstyle.font.color.rgb = color
        
        #adjust newstyle.paragraph_format parameters
        newstyle.paragraph_format.space_before = Pt(before_spacing)
        newstyle.paragraph_format.space_after = Pt(after_spacing)
        newstyle.paragraph_format.keep_together = keep_together
        if (not isinstance(line_spacing, (float, int))) or line_spacing < 1 :
            #raise valueError if invalid line_spacing
            #line_spacing = 0 makes text inapparent, although still present
            raise ValueError("Invalid line_space value, must be number > 1.0")
        else : 
            newstyle.paragraph_format.line_spacing = line_spacing
        #newstyle_format.left_indent = Inches(0.25)
        #newstyle_format.first_line_indent = Inches(-0.25)
        #newstyle_format.widow_control = True
        return newstyle

def remove_paragraphs(p, phrase = None) :
    """ 
    Args :
        p : a list of 'paragraphs' objects (example : docx.Document.paragraphs)
        phrase : string to be omitted
    Return :     
        
    """
    if phrase == None :
        raise ValueError("Input string to remove paragraphs containing that string")
    else : 
        for paragraph in p :
            #https://stackoverflow.com/questions/6930982/how-to-use-a-variable-inside-a-regular-expression
            #explanation on inserting variables into regular expressions
            if re.search(rf".*{re.escape(str(phrase))}.*", paragraph.text, re.IGNORECASE) :
                p.remove(paragraph)
        return p


def to_add () :
    ...
    
    

if __name__ == "__main__" :
    main()
    

